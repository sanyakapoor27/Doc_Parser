import streamlit as st
import pandas as pd
import io
import re
import base64

# Imports required for the Disaster Reporter App
import google.generativeai as genai
from io import BytesIO
import tempfile
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# import qn # Not strictly used in docx logic provided
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from dotenv import load_dotenv # Used for local testing of the provided code

# --- Configuration and Setup ---
# Set the page configuration for a wider layout (only one call is allowed in Streamlit)
st.set_page_config(layout="wide", page_title="Multi-Tool Data Dashboard")

# Define the threshold for unique values to trigger a multiselect dropdown
UNIQUE_VALUE_THRESHOLD = 50 

# Load environment variables (for Gemini API key)
load_dotenv()


# ==============================================================================
# 1. HELPER FUNCTIONS FOR DATA EXPLORER APP
# ==============================================================================

def load_data(uploaded_file):
    """Loads CSV or Excel data into a Pandas DataFrame."""
    try:
        if uploaded_file.name.endswith('.csv'):
            # Use io.TextIOWrapper to read CSV, allowing universal newline mode
            df = pd.read_csv(io.TextIOWrapper(uploaded_file, encoding='utf-8'))
        elif uploaded_file.name.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(uploaded_file)
        else:
            st.error("Unsupported file format. Please upload a CSV or Excel file.")
            return None
        return df
    except Exception as e:
        st.error(f"Error loading file: {e}")
        return None


# ==============================================================================
# 2. HELPER FUNCTIONS FOR DISASTER REPORTER APP
# ==============================================================================

def setup_gemini(api_key):
    """Initializes the Gemini client."""
    if not api_key:
        # Check for both environment variable (for local testing) and Streamlit secrets
        if "GEMINI_API" not in os.environ and "GEMINI_API" not in st.secrets:
            st.error("Gemini API key missing. Please set the GEMINI_API environment variable or Streamlit secret.")
            return False
    
    final_key = api_key or st.secrets.get("GEMINI_API", os.getenv("GEMINI_API"))

    try:
        genai.configure(api_key=final_key)
        return True
    except Exception as e:
        st.error(f"Error initializing Gemini: {e}")
        return False

def generate_docx_bytes(extracted_text: str, logo_path="logo.png") -> bytes:
    """
    Create a .docx in-memory with:
    - Logo on left
    - Title + Subtitle on right
    - Markdown table conversion
    """
    doc = Document()

    # ---------------------------
    # HEADER WITH LOGO + TITLES
    # ---------------------------
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # --- Left Cell: Logo ---
    left_cell = table.cell(0, 0)
    try:
        paragraph = left_cell.paragraphs[0]
        run = paragraph.add_run()
        # NOTE: This line requires 'logo.png' to be in the same directory.
        run.add_picture(logo_path, width=Inches(1.0)) 
    except Exception:
        left_cell.text = " [Logo Placeholder] " 

    # --- Right Cell: Headings ---
    right_cell = table.cell(0, 1)
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Title
    title_run = right_para.add_run("Disaster Situation Reporting\n")
    title_run.bold = True
    title_run.font.size = Pt(20)

    # Subtitle (smaller)
    subtitle_run = right_para.add_run("Disaster Management Centre\n")
    subtitle_run.font.size = Pt(12)
    subtitle_run.bold = False

    subtitle_run = right_para.add_run("Powered by IWMI")
    subtitle_run.font.size = Pt(12)
    subtitle_run.bold = False

    doc.add_paragraph("")  # spacing after header

    # ---------------------------
    # TEXT + TABLE PARSING
    # ---------------------------

    lines = extracted_text.splitlines()
    i = 0
    n = len(lines)

    def is_table_sep(line: str) -> bool:
        return bool(re.match(r'^\s*\|?\s*[:\- ]+(\s*\|\s*[:\- ]+)+\s*\|?\s*$', line))

    while i < n:
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # Detect markdown table
        if '|' in stripped and (i + 1 < n) and is_table_sep(lines[i + 1]):
            header_cells = [c.strip() for c in stripped.strip().strip('|').split('|')]
            i += 2

            table_rows = []
            while i < n:
                row_line = lines[i].strip()
                if not row_line or '|' not in row_line:
                    break
                cells = [c.strip() for c in row_line.strip().strip('|').split('|')]
                table_rows.append(cells)
                i += 1

            cols = len(header_cells)
            rows = 1 + len(table_rows)
            tbl = doc.add_table(rows=rows, cols=cols)
            tbl.style = "Table Grid"

            for ci, h in enumerate(header_cells):
                tbl.rows[0].cells[ci].text = h

            for r_idx, row in enumerate(table_rows, start=1):
                for ci, cell in enumerate(row):
                    if ci < cols:
                        tbl.rows[r_idx].cells[ci].text = cell

            doc.add_paragraph("")
            continue

        # Regular paragraph
        clean_line = stripped.replace("**", "").replace("##", "").replace("###", "").strip()
        if clean_line:
            doc.add_paragraph(clean_line)
        else:
            doc.add_paragraph("")

        i += 1

    # ---------------------------
    # RETURN BYTES
    # ---------------------------
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def run_extraction(file_handle):
    """Handles the document upload and Gemini API call for content extraction."""
    system_instruction = (
        "You are an expert document extraction and summarization agent. "
        "Extract ALL structured and unstructured information from the input document, "
        "preserving layout and tables in Markdown. "
        "Return three sections: "
        "## Content (English) — translated to English, "
        "## Content, "
        "## Summary — short English summary."
    )

    user_prompt = (
        "Analyze the uploaded document and extract all content. "
        "Output the three required sections using the correct headings."
    )

    uploaded_file_name = None
    temp_path = None
    try:
        with st.spinner("Uploading document…"):
            # 1. Save uploaded file to a temporary location
            temp_file = tempfile.NamedTemporaryFile(delete=False)
            file_handle.seek(0) # Ensure we read from the start of the file
            temp_file.write(file_handle.read())
            temp_path = temp_file.name
            temp_file.close() 
            
            # 2. Upload to Gemini using file path
            uploaded = genai.upload_file(
                path=temp_path,
                display_name=file_handle.name,
                mime_type=file_handle.type
            )
            uploaded_file_name = uploaded.name

        st.success(f"File uploaded: {file_handle.name}")

        contents = [user_prompt, uploaded]

        with st.spinner("Extracting content…"):
            model = genai.GenerativeModel('gemini-2.5-flash', system_instruction=system_instruction)
            response = model.generate_content(
                contents=contents
            )

        return response.text

    except Exception as e:
        st.error(f"Gemini API Error: {e}")
        return None

    finally:
        # Cleanup uploaded file from Gemini
        if uploaded_file_name:
            try:
                genai.delete_file(name=uploaded_file_name)
                # st.info("Temporary Gemini file cleaned.")
            except Exception as e:
                print("Gemini file cleanup failed:", e)

        # Cleanup temporary local file
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except:
                pass


# ==============================================================================
# 3. DATA EXPLORER APP (Original 'main' function)
# ==============================================================================

def data_explorer_app():

    """Application to upload, filter, and display data from CSV/Excel files."""

    # ===== HEADER SECTION =====
    with open("logo.png", "rb") as f:
        data = base64.b64encode(f.read()).decode()

    st.markdown(
        f"""
        <div style="text-align: center;">
            <img src="data:image/png;base64,{data}" style="width:70%; max-width:900px;">
        <p style="
                    font-style: italic; 
                    font-size:14px; 
                    margin-top:-5px;
                    margin-left:-143px;   /* adjust left shift */
                ">
                    Powered by International Water Management Institute (IWMI)
                </p>
        </div>
        """,
        unsafe_allow_html=True
    )

    st.markdown("<div style='margin-top:40px'></div>", unsafe_allow_html=True)

    st.write(" ### 📊 Data Filtering Tool")
    st.markdown("<div style='margin-top:20px'></div>", unsafe_allow_html=True)
    st.markdown("Upload your CSV or Excel file to begin filtering and exploring the data.")

    st.markdown("<div style='margin-top:30px'></div>", unsafe_allow_html=True)

    # 1. File Uploader
    uploaded_file = st.file_uploader(
        "Choose a CSV or Excel File",
        type=['csv', 'xlsx', 'xls'],
        help="Upload any structured data file to instantly get filtering tools."
    )

    # Make upload button smaller
    st.markdown("""
        <style>

            /* --- ADJUST H2 "Upload Document" TITLE SPACING --- */
            h2 {
                margin-top: 0px !important;
                margin-bottom: 6px !important;
            }

            /* --- FILE UPLOADER GENERAL RESET --- */
            div[data-testid="stFileUploader"] {
                margin-top: 0 !important;
                margin-bottom: 0 !important;
                padding: 0 !important;
            }

            /* --- INNER WRAPPER (holds the dropzone) --- */
            div[data-testid="stFileUploader"] > div {
                margin: 0 !important;
                padding: 0 !important;
            }

            /* --- DROPZONE CONTAINER --- */
            [data-testid="stFileUploaderDropzone"] {
                width: 400px !important;  
                padding: 6px 10px !important;      /* smaller inner space */
                margin-top: -2px !important;       /* pull closer to title */
                border-radius: 6px !important;
            }

            /* --- DROPZONE TEXT + ICON WRAPPER --- */
            [data-testid="stFileUploaderDropzone"] div {
                margin: -4px 0 -2px 0 !important;  /* tighten inner vertical spacing */
            }

            /* --- BUTTON INSIDE DROPZONE --- */
            [data-testid="stFileUploaderDropzone"] button {
                padding: 3px 10px !important;      /* smaller button */
                font-size: 0.85rem !important;
                margin: 0 !important;
            }

            /* --- LIMIT SIZE TEXT (small gray text) --- */
            [data-testid="stFileUploaderDropzone"] small {
                margin-top: -8px !important;       /* reduces white-space below icon */
                display: block !important;
            }

            /* --- GENERAL BLOCK SPACING FIX (stVerticalBlock) --- */
            div[data-testid="stVerticalBlock"] {
                gap: 0px !important;
                padding: 0 !important;
                margin: 0 !important;
            }

        </style>
        """, unsafe_allow_html=True)
    
    st.markdown("<div style='margin-top:40px'></div>", unsafe_allow_html=True)

    if uploaded_file is None:
        st.info("Awaiting file upload...")
        return

    # Load data
    df = load_data(uploaded_file)
    if df is None:
        return

    # Initialize the filtered DataFrame with the original data
    df_filtered = df.copy()

    # --- 2. Dynamic Filtering Section ---
    st.sidebar.markdown("---")
    st.sidebar.header("Active Data Filters")
    st.sidebar.markdown(f"**Loaded:** `{uploaded_file.name}`")
    st.sidebar.markdown(f"**Rows:** {len(df)} | **Columns:** {len(df.columns)}")
    st.sidebar.markdown("---")

    st.subheader("Filter Settings")

    # Dynamic column filtering containers (using columns for better layout)
    num_cols_to_display = 4
    filter_columns = st.columns(num_cols_to_display)
    
    col_index = 0
    
    # Iterate through all columns in the DataFrame
    for column in df.columns:
        # Place the filter widget in the next available column container
        with filter_columns[col_index % num_cols_to_display]:
            
            # Identify unique values and count
            unique_values = df[column].dropna().astype(str).unique()
            n_unique = len(unique_values)

            # Check if the column is suitable for a multi-select dropdown
            if n_unique <= UNIQUE_VALUE_THRESHOLD and n_unique > 0:
                # 3. Dropdown/Multi-select for low unique values
                st.caption(f"Categorical Filter ({n_unique} unique)")
                selected_values = st.multiselect(
                    f"Select values for '{column}'",
                    options=unique_values,
                    default=[],
                    key=f"select_{column}"
                )

                # Apply filter: if selections are made, filter the DataFrame
                if selected_values:
                    # Convert column data to string for consistent comparison
                    df_filtered = df_filtered[df_filtered[column].astype(str).isin(selected_values)]
            
            else:
                # 2. Text Input for high unique values or general search
                st.caption("Search Filter (Use regex, e.g., `^A` or `.*ing`)")
                search_term = st.text_input(
                    f"Search in '{column}'",
                    placeholder="Enter search string or regex...",
                    key=f"search_{column}"
                )

                # Apply filter: if a search term is provided, filter the DataFrame
                if search_term:
                    try:
                        # Use a case-insensitive regex search
                        df_filtered = df_filtered[
                            df_filtered[column].astype(str).str.contains(search_term, case=False, na=False, regex=True)
                        ]
                    except re.error:
                        st.error(f"Invalid Regular Expression used in '{column}' filter.")
                        # Stop filtering if regex is invalid
                        return

        col_index += 1
        
    st.markdown("---")


    # --- 4. Dashboard Display ---
    st.subheader("Filtered Data Dashboard")
    
    # Display statistics
    st.metric(
        label="Total Rows Remaining", 
        value=f"{len(df_filtered):,}",
        delta=f"- {len(df) - len(df_filtered):,} rows filtered out",
        delta_color="inverse"
    )

    if len(df_filtered) == 0:
        st.warning("No data matches the current filter criteria.")
    else:
        # Display the filtered data table
        st.dataframe(df_filtered, width='stretch', height=500)

        # Optional: Allow download of the filtered data
        csv_export = df_filtered.to_csv(index=False).encode('utf-8')
        st.download_button(
            label="Download Filtered Data as CSV",
            data=csv_export,
            file_name='filtered_data.csv',
            mime='text/csv',
            help="Click here to download the currently filtered dataset."
        )


# ==============================================================================
# 4. DISASTER REPORTER APP
# ==============================================================================

def disaster_reporter_app():
    """Application for document extraction and reporting using Gemini."""
    
    # ===== HEADER SECTION =====
    with open("logo.png", "rb") as f:
        data = base64.b64encode(f.read()).decode()

    st.markdown(
        f"""
        <div style="text-align: center;">
            <img src="data:image/png;base64,{data}" style="width:70%; max-width:900px;">
        <p style="
                    font-style: italic; 
                    font-size:14px; 
                    margin-top:-5px;
                    margin-left:-143px;   /* adjust left shift */
                ">
                    Powered by International Water Management Institute (IWMI)
                </p>
        </div>
        """,
        unsafe_allow_html=True
    )

    # Make spacing tighter
    st.markdown("<div style='margin-top:40px'></div>", unsafe_allow_html=True)

    # ===== Upload Section =====
    st.write("### 📄 Upload Document")

    uploaded_file = st.file_uploader(
        "",
        type=["pdf", "jpg", "jpeg", "png"],
        help="Upload a PDF or image for extraction."
    )

    # Reduce extra spacing after uploader
    st.markdown("<div style='margin-top:-25px'></div>", unsafe_allow_html=True)

    # Make upload button smaller
    st.markdown("""
        <style>

            /* --- ADJUST H2 "Upload Document" TITLE SPACING --- */
            h2 {
                margin-top: 0px !important;
                margin-bottom: 6px !important;
            }

            /* --- FILE UPLOADER GENERAL RESET --- */
            div[data-testid="stFileUploader"] {
                margin-top: 0 !important;
                margin-bottom: 0 !important;
                padding: 0 !important;
            }

            /* --- INNER WRAPPER (holds the dropzone) --- */
            div[data-testid="stFileUploader"] > div {
                margin: 0 !important;
                padding: 0 !important;
            }

            /* --- DROPZONE CONTAINER --- */
            [data-testid="stFileUploaderDropzone"] {
                width: 400px !important;  
                padding: 6px 10px !important;      /* smaller inner space */
                margin-top: -2px !important;       /* pull closer to title */
                border-radius: 6px !important;
            }

            /* --- DROPZONE TEXT + ICON WRAPPER --- */
            [data-testid="stFileUploaderDropzone"] div {
                margin: -4px 0 -2px 0 !important;  /* tighten inner vertical spacing */
            }

            /* --- BUTTON INSIDE DROPZONE --- */
            [data-testid="stFileUploaderDropzone"] button {
                padding: 3px 10px !important;      /* smaller button */
                font-size: 0.85rem !important;
                margin: 0 !important;
            }

            /* --- LIMIT SIZE TEXT (small gray text) --- */
            [data-testid="stFileUploaderDropzone"] small {
                margin-top: -8px !important;       /* reduces white-space below icon */
                display: block !important;
            }

            /* --- GENERAL BLOCK SPACING FIX (stVerticalBlock) --- */
            div[data-testid="stVerticalBlock"] {
                gap: 0px !important;
                padding: 0 !important;
                margin: 0 !important;
            }

        </style>
        """, unsafe_allow_html=True)
    
    st.markdown("<div style='margin-top:25px'></div>", unsafe_allow_html=True)

    st.markdown("### Instructions")
    st.markdown(""" 
            1. Upload PDF or Image 
            2. Click Extract 
            3. Download the generated DOCX report.
    """)

    # Use os.getenv for flexibility in environment setup
    api_key = os.getenv("GEMINI_API") or st.secrets.get("GEMINI_API")

    # Always show button after upload
    if uploaded_file:
        # Setup Gemini once
        gemini_ok = setup_gemini(api_key)

        if gemini_ok and st.button("🚀 Extract & Summarize Document", type="primary"):
            # --- 1. EXTRACTION LOGIC (Only runs when button is clicked) ---
            extracted_content = run_extraction(uploaded_file)
            
            # Store results in session state
            if extracted_content:
                st.session_state["full_output"] = extracted_content
                st.session_state["has_results"] = True
            else:
                st.session_state["full_output"] = ""
                st.session_state["has_results"] = False


    # --- 2. DISPLAY AND DOWNLOAD LOGIC (Runs on every rerun if results exist) ---
    if st.session_state.get("has_results", False):
        full_output = st.session_state.get("full_output", "")
        st.header("✅ Extraction Results")

        # Regular expression to find the start of the Summary section
        pattern = r"\s*##\s*Summary"
        match = re.search(pattern, full_output, re.IGNORECASE)

        content_for_pdf = full_output
        summary_section = ""

        if match:
            idx = match.start()
            # Content is everything BEFORE the Summary heading
            content_for_pdf = full_output[:idx].strip()
            # Summary is the Summary heading and everything AFTER it
            summary_section = full_output[idx:].strip()
            
            # Display content
            st.markdown(content_for_pdf)

            st.markdown("---")
            st.subheader("Summary")
            # Display summary, removing the heading
            st.markdown(summary_section.replace("## Summary", "", 1).strip())
            st.markdown("---")
        else:
            # Display full output if no summary heading found
            st.markdown(full_output)
            st.warning("Summary section not found — exporting whole output to DOCX.")

        # DOCX Download
        try:
            # Pass only the non-summary content for DOCX generation
            docx_bytes = generate_docx_bytes(content_for_pdf, "logo.png")
            st.download_button(
                label="⬇️ Download Extracted Content (DOCX)",
                data=docx_bytes,
                file_name="extracted_document_content.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"DOCX generation error: {e}")


# ==============================================================================
# 5. MAIN ROUTER
# ==============================================================================

def main_router():
    """Routes the user to the selected Streamlit application."""
    
    with st.sidebar:
        st.title("Multi-Tool Dashboard")
        st.markdown("---")
        app_mode = st.radio(
            "Choose your tool:",
            ["Data Filtering", "Disaster Situation Reporting"],
            index=0 # Default to Data Explorer
        )

    # Run the selected application
    if app_mode == "Data Filtering":
        data_explorer_app()
    elif app_mode == "Disaster Situation Reporting":
        disaster_reporter_app()

if __name__ == "__main__":
    main_router()